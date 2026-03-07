"""
Generate GMAT Study Plan Unit Test Cases (.docx)
Produces: GMAT_Study_Plan_Test_Cases.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date

# ── Helpers ──────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(s)

def add_test_table(doc, rows):
    """Add a formatted test-case table.  rows = list of dicts with keys:
       id, case, precondition, action, expected
    """
    headers = ['ID', 'Test Case', 'Precondition (xAPI State)', 'Action', 'Expected Result', 'Status']
    widths  = [Cm(1.2), Cm(4.5), Cm(4.5), Cm(3), Cm(4.5), Cm(1.5)]

    tbl = doc.add_table(rows=1 + len(rows), cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '00409E')

    # Data rows
    for ri, row in enumerate(rows):
        vals = [row.get('id',''), row.get('case',''), row.get('precondition',''),
                row.get('action',''), row.get('expected',''), '']
        for ci, v in enumerate(vals):
            cell = tbl.rows[ri+1].cells[ci]
            cell.text = v
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
        # Zebra striping
        if ri % 2 == 1:
            for ci in range(6):
                set_cell_shading(tbl.rows[ri+1].cells[ci], 'EEF3FB')

    # Column widths
    for ri_idx in range(len(tbl.rows)):
        for ci_idx, w in enumerate(widths):
            tbl.rows[ri_idx].cells[ci_idx].width = w

    doc.add_paragraph('')  # spacer


# ── Build Document ───────────────────────────────────────────────────

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── Cover Page ───────────────────────────────────────────────────────

for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('GMAT Study Plan')
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x00, 0x40, 0x9E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Unit Test Cases')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x00, 0x2B, 0x6B)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Gurutor Staging - Manual QA Document')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Version 1.0  |  {date.today().strftime("%B %d, %Y")}')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ── Table of Contents ────────────────────────────────────────────────

doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Section A - Plan Structure Tests',
    'Section B - xAPI Signal Functions',
    'Section C - Verbal Suggestion Logic (Verbal-First)',
    'Section D - Quant Suggestion Logic (Verbal-First)',
    'Section E - Data Insights Suggestion Logic',
    'Section F - Quant-First Path Differences',
    'Section G - Edge Cases & Data Integrity',
    'Section H - Suggest Box Rendering',
]
for i, item in enumerate(toc_items, 1):
    doc.add_paragraph(f'{i}.  {item}', style='List Number')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION A - Plan Structure Tests
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Section A - Plan Structure Tests', level=1)
doc.add_paragraph(
    'Verify that each learning path builds with the correct lessons '
    'in the correct learn / practice / review slots for every unit.'
)

# -- A1 Verbal-First --
doc.add_heading('A1. Verbal-First Path', level=2)

vf_units = [
    # (unit label, learn, practice, review)
    ('Verbal Unit 1: Foundations', 'intro_verbal, intro_quant, intro_di, cr_lesson_1, cr_lesson_2, rc_lesson_1', 'cr_exercise_1, cr_exercise_2', '(none)'),
    ('Verbal Unit 2: Core Skills', 'cr_lesson_3, cr_lesson_4, rc_lesson_2, rc_lesson_3', 'cr_exercise_3, rc_exercise_1', '(none)'),
    ('Verbal Unit 3: Strategic CR', 'cr_lesson_5', 'cr_exercise_4', 'verbal_review_2'),
    ('Verbal Unit 4: Regular Arguments', 'cr_lesson_6', 'cr_exercise_5', 'verbal_review_3'),
    ('Verbal Unit 5: Explanation CR', 'cr_lesson_7', 'cr_exercise_6', 'verbal_review_4'),
    ('Verbal Unit 6: Advanced CR', 'cr_lesson_8, cr_lesson_9', 'cr_exercise_7, cr_exercise_8', 'verbal_review_5'),
    ('Quant Unit 1: Intro', 'intro_quant', '(none)', '(none)'),
    ('Quant Unit 2: Algebra & Translation', 'pss_lesson_1, algebra_1, word_problems_1, number_props_1', 'quant_exercise_1', '(none)'),
    ('Quant Unit 3: Estimation & Multi-Step', 'pss_lesson_2, number_props_2, algebra_2, word_problems_2, fprs_1', 'quant_exercise_2', 'quant_review_2'),
    ('Quant Unit 4: Advanced WP', 'fprs_2, algebra_3, word_problems_3, word_problems_4', 'quant_exercise_3', 'quant_review_3'),
    ('Quant Unit 5: Systems & Probability', 'algebra_4, word_problems_5, word_problems_6', 'quant_exercise_4', 'quant_review_4'),
    ('Quant Unit 6: Patterns & Edge Cases', 'number_props_3, word_problems_7, algebra_5', 'quant_exercise_5', 'quant_review_5'),
    ('DI Unit 1: Data Sufficiency', 'intro_di, di_lesson_1, di_lesson_2, di_lesson_3', '(none)', 'quant_review_6'),
    ('DI Unit 2: Visual & Tabular', 'di_lesson_4, di_lesson_5', '(none)', '(none)'),
    ('DI Unit 3: Multi-Source', 'di_lesson_6, di_lesson_7', '(none)', '(none)'),
]

rows_a1 = []
for i, (label, learn, practice, review) in enumerate(vf_units, 1):
    rows_a1.append({
        'id': f'A1.{i}',
        'case': f'{label} — structure',
        'precondition': 'User preference = verbal',
        'action': 'Load study plan page',
        'expected': f'Learn: {learn}\nPractice: {practice}\nReview: {review}',
    })
add_test_table(doc, rows_a1)

# -- A2 Quant-First --
doc.add_heading('A2. Quant-First Path', level=2)

qf_units = [
    ('Quant Unit 1: Intro', 'intro_quant', '(none)', '(none)'),
    ('Quant Unit 2: Algebra & Translation', 'pss_lesson_1, algebra_1, word_problems_1, number_props_1', 'quant_exercise_1', '(none)'),
    ('Quant Unit 3: Estimation', 'pss_lesson_2, number_props_2, algebra_2, word_problems_2, fprs_1', 'quant_exercise_2', 'quant_review_2'),
    ('Quant Unit 4: Advanced WP', 'fprs_2, algebra_3, word_problems_3, word_problems_4', 'quant_exercise_3', 'quant_review_3'),
    ('Quant Unit 5: Systems', 'algebra_4, word_problems_5, word_problems_6', 'quant_exercise_4', 'quant_review_4'),
    ('Quant Unit 6: Patterns', 'number_props_3, word_problems_7, algebra_5', 'quant_exercise_5', 'quant_review_5'),
    ('Verbal Unit 1: Foundations', 'intro_verbal, cr_lesson_1, cr_lesson_2, rc_lesson_1', 'cr_exercise_1, cr_exercise_2', 'quant_review_6'),
    ('Verbal Unit 2: Core Skills', 'cr_lesson_3, cr_lesson_4, rc_lesson_2, rc_lesson_3', 'cr_exercise_3, rc_exercise_1', '(none)'),
    ('Verbal Unit 3: Strategic CR', 'cr_lesson_5', 'cr_exercise_4', 'verbal_review_2'),
    ('Verbal Unit 4: Regular Arguments', 'cr_lesson_6', 'cr_exercise_5', 'verbal_review_3'),
    ('Verbal Unit 5: Explanation CR', 'cr_lesson_7', 'cr_exercise_6', 'verbal_review_4'),
    ('Verbal Unit 6: Advanced CR', 'cr_lesson_8, cr_lesson_9', 'cr_exercise_7, cr_exercise_8', 'verbal_review_5'),
    ('DI Unit 1: Data Sufficiency', 'intro_di, di_lesson_1, di_lesson_2, di_lesson_3', '(none)', '(none)'),
    ('DI Unit 2: Visual & Tabular', 'di_lesson_4, di_lesson_5', '(none)', '(none)'),
    ('DI Unit 3: Multi-Source', 'di_lesson_6, di_lesson_7', '(none)', '(none)'),
]

rows_a2 = []
for i, (label, learn, practice, review) in enumerate(qf_units, 1):
    rows_a2.append({
        'id': f'A2.{i}',
        'case': f'{label} — structure',
        'precondition': 'User preference = quant',
        'action': 'Load study plan page',
        'expected': f'Learn: {learn}\nPractice: {practice}\nReview: {review}',
    })
add_test_table(doc, rows_a2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION B - xAPI Signal Functions
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Section B - xAPI Signal Functions', level=1)
doc.add_paragraph(
    'Test the 3-state result functions that parse xAPI pass/fail signals. '
    'Every function returns exactly one of: "fail" | "pass" | "none".'
)

# B1 - Exercise Result
doc.add_heading('B1. gmat_sp_get_exercise_result()', level=2)
doc.add_paragraph('Returns pass/fail/none for individual CR exercises (4-8).')

cr_exercises = [
    ('cr_exercise_4', 'CR_Exercise_4_Pass_or_Fail', 'Plan Arguments'),
    ('cr_exercise_5', 'CR_Exercise_5_Pass_or_Fail', 'Regular Arguments'),
    ('cr_exercise_6', 'CR_Exercise_6_Pass_or_Fail', 'Explanation Arguments'),
    ('cr_exercise_7', 'CR_Exercise_7_Pass_or_Fail', 'Describe the Role'),
    ('cr_exercise_8', 'CR_Exercise_8_Pass_or_Fail', 'Inference & Discrepancy'),
]

rows_b1 = []
for i, (key, var, topic) in enumerate(cr_exercises):
    base = i * 3
    rows_b1.append({
        'id': f'B1.{base+1}', 'case': f'{key} — Fail',
        'precondition': f'{var} = "Fail"',
        'action': f'Call gmat_sp_get_exercise_result(user, "{key}")',
        'expected': 'Returns "fail"',
    })
    rows_b1.append({
        'id': f'B1.{base+2}', 'case': f'{key} — Pass',
        'precondition': f'{var} = "Pass"',
        'action': f'Call gmat_sp_get_exercise_result(user, "{key}")',
        'expected': 'Returns "pass"',
    })
    rows_b1.append({
        'id': f'B1.{base+3}', 'case': f'{key} — No signal',
        'precondition': f'No {var} in xAPI data',
        'action': f'Call gmat_sp_get_exercise_result(user, "{key}")',
        'expected': 'Returns "none"',
    })
add_test_table(doc, rows_b1)

# B2 - Review Result
doc.add_heading('B2. gmat_sp_get_review_result()', level=2)
doc.add_paragraph(
    'Returns pass/fail/none for review sets. Multi-variable aggregation: '
    'if ANY variable mapped to the review shows "Fail", the result is "fail".'
)

reviews = [
    ('verbal_review_2', 'Unit2_Verbal_Review_Pass_or_Fail'),
    ('verbal_review_3', 'Unit3_Verbal_Review_Pass_or_Fail'),
    ('verbal_review_4', 'Unit4_Verbal_Review_Pass_or_Fail'),
    ('verbal_review_5', 'Unit5_Verbal_Review_Pass_or_Fail'),
    ('quant_review_2', 'QRS_Unit2_* variables'),
    ('quant_review_3', 'QRS_Unit3_* variables'),
    ('quant_review_4', 'QRS_Unit4_* variables'),
    ('quant_review_5', 'QRS_Unit5_* variables'),
]

rows_b2 = []
for i, (key, var) in enumerate(reviews):
    base = i * 3
    rows_b2.append({
        'id': f'B2.{base+1}', 'case': f'{key} — Fail',
        'precondition': f'{var} includes at least one "Fail"',
        'action': f'Call gmat_sp_get_review_result(user, "{key}")',
        'expected': 'Returns "fail"',
    })
    rows_b2.append({
        'id': f'B2.{base+2}', 'case': f'{key} — Pass',
        'precondition': f'{var} all show "Pass"',
        'action': f'Call gmat_sp_get_review_result(user, "{key}")',
        'expected': 'Returns "pass"',
    })
    rows_b2.append({
        'id': f'B2.{base+3}', 'case': f'{key} — None',
        'precondition': f'No {var} in xAPI data',
        'action': f'Call gmat_sp_get_review_result(user, "{key}")',
        'expected': 'Returns "none"',
    })

# Extra: multi-variable mixed case
rows_b2.append({
    'id': 'B2.25', 'case': 'quant_review_2 — Mixed (1 fail, 3 pass)',
    'precondition': 'QRS_Unit2_ALG1 = Pass, QRS_Unit2_NP1 = Fail, QRS_Unit2_WP1 = Pass, QRS_Unit2_PSS1 = Pass',
    'action': 'Call gmat_sp_get_review_result(user, "quant_review_2")',
    'expected': 'Returns "fail" (ANY fail = fail)',
})
add_test_table(doc, rows_b2)

# B3 - Quant Exercise Failures
doc.add_heading('B3. gmat_sp_get_quant_exercise_failures()', level=2)
doc.add_paragraph(
    'Returns array of lesson keys that failed within a specific quant exercise (QLE_* signals).'
)

qle_data = [
    (1, ['algebra_1', 'number_props_1', 'word_problems_1', 'pss_lesson_1']),
    (2, ['algebra_2', 'number_props_2', 'word_problems_2', 'pss_lesson_2', 'fprs_1']),
    (3, ['algebra_3', 'fprs_2', 'word_problems_3', 'word_problems_4']),
    (4, ['algebra_4', 'word_problems_5', 'word_problems_6']),
    (5, ['algebra_5', 'number_props_3', 'word_problems_7']),
]

rows_b3 = []
tc = 1
for ex_num, lessons in qle_data:
    rows_b3.append({
        'id': f'B3.{tc}', 'case': f'Exercise {ex_num} — All fail',
        'precondition': f'All QLE_{ex_num}_* = "Fail"',
        'action': f'Call gmat_sp_get_quant_exercise_failures(user, {ex_num}, ...)',
        'expected': f'Returns [{", ".join(lessons)}]',
    })
    tc += 1
    rows_b3.append({
        'id': f'B3.{tc}', 'case': f'Exercise {ex_num} — All pass',
        'precondition': f'All QLE_{ex_num}_* = "Pass"',
        'action': f'Call gmat_sp_get_quant_exercise_failures(user, {ex_num}, ...)',
        'expected': 'Returns [] (empty)',
    })
    tc += 1
    rows_b3.append({
        'id': f'B3.{tc}', 'case': f'Exercise {ex_num} — Partial fail',
        'precondition': f'QLE_{ex_num}_{lessons[0].upper().split("_")[0][:3]}* = Fail, rest = Pass',
        'action': f'Call gmat_sp_get_quant_exercise_failures(user, {ex_num}, ...)',
        'expected': f'Returns [{lessons[0]}]',
    })
    tc += 1
    rows_b3.append({
        'id': f'B3.{tc}', 'case': f'Exercise {ex_num} — No signals',
        'precondition': f'No QLE_{ex_num}_* variables',
        'action': f'Call gmat_sp_get_quant_exercise_failures(user, {ex_num}, ...)',
        'expected': 'Returns [] (empty)',
    })
    tc += 1
add_test_table(doc, rows_b3)

# B4 - Learn Lesson Failures
doc.add_heading('B4. gmat_sp_get_learn_lesson_failures()', level=2)
doc.add_paragraph(
    'Returns array of learn lesson keys that have explicit "Fail" signals '
    '(direct *_Lesson_* variables only, excludes QLE/QRS).'
)

rows_b4 = [
    {'id': 'B4.1', 'case': 'Single lesson failure',
     'precondition': 'ALG_Lesson_1_Pass_or_Fail = "Fail", rest = "Pass"',
     'action': 'Call with learn_keys = [pss_lesson_1, algebra_1, word_problems_1, number_props_1]',
     'expected': 'Returns [algebra_1]'},
    {'id': 'B4.2', 'case': 'Multiple lesson failures',
     'precondition': 'ALG_Lesson_1 = "Fail", WP_Lesson_1 = "Fail"',
     'action': 'Call with learn_keys = [pss_lesson_1, algebra_1, word_problems_1, number_props_1]',
     'expected': 'Returns [algebra_1, word_problems_1]'},
    {'id': 'B4.3', 'case': 'All lessons pass',
     'precondition': 'All *_Lesson_* variables = "Pass"',
     'action': 'Call with any learn_keys array',
     'expected': 'Returns [] (empty)'},
    {'id': 'B4.4', 'case': 'No lesson signals',
     'precondition': 'No *_Lesson_* variables in xAPI data',
     'action': 'Call with any learn_keys array',
     'expected': 'Returns [] (empty)'},
    {'id': 'B4.5', 'case': 'QLE failure not included',
     'precondition': 'QLE_1_ALG1 = "Fail" (but NO ALG_Lesson_1 signal)',
     'action': 'Call with learn_keys = [algebra_1]',
     'expected': 'Returns [] (QLE signals excluded)'},
]
add_test_table(doc, rows_b4)

# B5 - Quant Review Failures
doc.add_heading('B5. gmat_sp_get_quant_review_failures()', level=2)
doc.add_paragraph(
    'Returns lesson keys that failed within a quant review set (QRS_Unit* signals). '
    'Uses gmat_sp_get_qrs_lesson_map() to resolve topic suffixes to lesson keys.'
)

qrs_data = [
    (2, {'ALG1': 'algebra_1', 'NP1': 'number_props_1', 'WP1': 'word_problems_1', 'PSS1': 'pss_lesson_1'}),
    (3, {'ALG2': 'algebra_2', 'NP2': 'number_props_2', 'WP2': 'word_problems_2', 'PSS2': 'pss_lesson_2', 'FPR1': 'fprs_1'}),
    (4, {'ALG3': 'algebra_3', 'WP3': 'word_problems_3', 'WP4': 'word_problems_4', 'FRP4': 'fprs_2'}),
    (5, {'ALG4': 'algebra_4', 'WP5': 'word_problems_5', 'WP6': 'word_problems_6'}),
]

rows_b5 = []
tc = 1
for unit, topics in qrs_data:
    all_lessons = list(topics.values())
    rows_b5.append({
        'id': f'B5.{tc}', 'case': f'QRS Unit {unit} — All fail',
        'precondition': f'All QRS_Unit{unit}_* = "Fail"',
        'action': f'Call gmat_sp_get_quant_review_failures(user, {unit})',
        'expected': f'Returns [{", ".join(all_lessons)}]',
    })
    tc += 1
    first_topic = list(topics.keys())[0]
    rows_b5.append({
        'id': f'B5.{tc}', 'case': f'QRS Unit {unit} — Partial fail ({first_topic})',
        'precondition': f'QRS_Unit{unit}_{first_topic} = "Fail", rest = "Pass"',
        'action': f'Call gmat_sp_get_quant_review_failures(user, {unit})',
        'expected': f'Returns [{topics[first_topic]}]',
    })
    tc += 1
    rows_b5.append({
        'id': f'B5.{tc}', 'case': f'QRS Unit {unit} — All pass',
        'precondition': f'All QRS_Unit{unit}_* = "Pass"',
        'action': f'Call gmat_sp_get_quant_review_failures(user, {unit})',
        'expected': 'Returns [] (empty)',
    })
    tc += 1
    rows_b5.append({
        'id': f'B5.{tc}', 'case': f'QRS Unit {unit} — No signals',
        'precondition': f'No QRS_Unit{unit}_* in xAPI data',
        'action': f'Call gmat_sp_get_quant_review_failures(user, {unit})',
        'expected': 'Returns [] (empty)',
    })
    tc += 1
add_test_table(doc, rows_b5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION C - Verbal Suggestion Logic (Verbal-First)
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Section C - Verbal Suggestion Logic (Verbal-First)', level=1)
doc.add_paragraph(
    'Test suggestion triggers in the verbal section when user preference is "verbal" (verbal-first path). '
    'Verbal Units 1-3 have NO suggestion logic. Units 4-6 have review suggestion boxes '
    'triggered by previous unit CR exercise failures.'
)

SUGGEST_TEXT = 'Before completing the Review Set, we suggest that you revisit the following lessons to reinforce the skills that the Review Set tests.'

rows_c = [
    # Units 1-2: no suggestions
    {'id': 'C.1', 'case': 'Verbal U1 — No suggestion boxes',
     'precondition': 'Any xAPI state', 'action': 'Load verbal-first study plan',
     'expected': 'Verbal Unit 1 has NO practice suggest box and NO review suggest box'},
    {'id': 'C.2', 'case': 'Verbal U2 — No suggestion boxes',
     'precondition': 'Any xAPI state', 'action': 'Load verbal-first study plan',
     'expected': 'Verbal Unit 2 has NO practice suggest box and NO review suggest box'},
    {'id': 'C.3', 'case': 'Verbal U3 — No suggestion boxes',
     'precondition': 'Any xAPI state', 'action': 'Load verbal-first study plan',
     'expected': 'Verbal Unit 3 has NO review suggest box (review_suggest is empty)'},

    # Unit 4: review suggest when CR Exercise 4 fails
    {'id': 'C.4', 'case': 'Verbal U4 — CR Ex 4 FAIL',
     'precondition': 'CR_Exercise_4_Pass_or_Fail = "Fail"',
     'action': 'Load verbal-first study plan, scroll to Verbal Unit 4 Review',
     'expected': f'Orange suggest box appears ABOVE verbal_review_3.\nText: "{SUGGEST_TEXT}"\nRedo link: cr_lesson_5 (CR Plan Arguments)'},
    {'id': 'C.5', 'case': 'Verbal U4 — CR Ex 4 PASS',
     'precondition': 'CR_Exercise_4_Pass_or_Fail = "Pass"',
     'action': 'Load verbal-first study plan, scroll to Verbal Unit 4 Review',
     'expected': 'NO review suggest box appears'},
    {'id': 'C.6', 'case': 'Verbal U4 — CR Ex 4 NO SIGNAL',
     'precondition': 'No CR_Exercise_4_Pass_or_Fail in xAPI',
     'action': 'Load verbal-first study plan, scroll to Verbal Unit 4 Review',
     'expected': 'NO review suggest box appears'},

    # Unit 5: review suggest when CR Exercise 5 fails
    {'id': 'C.7', 'case': 'Verbal U5 — CR Ex 5 FAIL',
     'precondition': 'CR_Exercise_5_Pass_or_Fail = "Fail"',
     'action': 'Load verbal-first study plan, scroll to Verbal Unit 5 Review',
     'expected': f'Orange suggest box appears ABOVE verbal_review_4.\nText: "{SUGGEST_TEXT}"\nRedo link: cr_lesson_6 (CR Regular Arguments)'},
    {'id': 'C.8', 'case': 'Verbal U5 — CR Ex 5 PASS',
     'precondition': 'CR_Exercise_5_Pass_or_Fail = "Pass"',
     'action': 'Load study plan', 'expected': 'NO review suggest box at Verbal U5'},
    {'id': 'C.9', 'case': 'Verbal U5 — CR Ex 5 NO SIGNAL',
     'precondition': 'No CR_Exercise_5_Pass_or_Fail',
     'action': 'Load study plan', 'expected': 'NO review suggest box at Verbal U5'},

    # Unit 6: review suggest when CR Exercise 6 fails
    {'id': 'C.10', 'case': 'Verbal U6 — CR Ex 6 FAIL',
     'precondition': 'CR_Exercise_6_Pass_or_Fail = "Fail"',
     'action': 'Load verbal-first study plan, scroll to Verbal Unit 6 Review',
     'expected': f'Orange suggest box appears ABOVE verbal_review_5.\nText: "{SUGGEST_TEXT}"\nRedo link: cr_lesson_7 (CR Explanation Arguments)'},
    {'id': 'C.11', 'case': 'Verbal U6 — CR Ex 6 PASS',
     'precondition': 'CR_Exercise_6_Pass_or_Fail = "Pass"',
     'action': 'Load study plan', 'expected': 'NO review suggest box at Verbal U6'},
    {'id': 'C.12', 'case': 'Verbal U6 — CR Ex 6 NO SIGNAL',
     'precondition': 'No CR_Exercise_6_Pass_or_Fail',
     'action': 'Load study plan', 'expected': 'NO review suggest box at Verbal U6'},
]
add_test_table(doc, rows_c)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION D - Quant Suggestion Logic (Verbal-First)
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Section D - Quant Suggestion Logic (Verbal-First)', level=1)
doc.add_paragraph(
    'Test all three suggestion types in the quant section (verbal-first path):\n'
    '1. Practice suggest — QLE failures + direct lesson failures for current unit\n'
    '2. Review suggest — QLE failures from previous unit\n'
    '3. Cross-section suggest — Verbal review pass/fail/none from corresponding verbal unit'
)

PRACTICE_SUGGEST = 'Before completing the Quant Exercise, we suggest that you revisit the following lessons to reinforce the skills that the Exercise tests.'
REVIEW_SUGGEST = 'Before completing the Review Set, we suggest that you revisit the following lessons to reinforce the skills that the Review Set tests.'

rows_d = []

# D1 - Quant Unit 2 (practice suggest only)
rows_d.append({'id': 'D.1', 'case': 'Q-U2 Practice — QLE_1 failures',
    'precondition': 'QLE_1_ALG1 = Fail, QLE_1_WP1 = Fail, rest = Pass',
    'action': 'Load study plan, Quant Unit 2 Practice',
    'expected': f'Practice suggest box with redo: algebra_1, word_problems_1.\nText: "{PRACTICE_SUGGEST}"'})
rows_d.append({'id': 'D.2', 'case': 'Q-U2 Practice — Direct lesson failures',
    'precondition': 'ALG_Lesson_1_Pass_or_Fail = Fail (no QLE signals)',
    'action': 'Load study plan, Quant Unit 2 Practice',
    'expected': 'Practice suggest box with redo: algebra_1'})
rows_d.append({'id': 'D.3', 'case': 'Q-U2 Practice — All pass',
    'precondition': 'All QLE_1_* = Pass, all direct lessons = Pass',
    'action': 'Load study plan', 'expected': 'NO practice suggest box at Quant U2'})
rows_d.append({'id': 'D.4', 'case': 'Q-U2 Practice — No signals',
    'precondition': 'No QLE_1_* or direct lesson signals',
    'action': 'Load study plan', 'expected': 'NO practice suggest box at Quant U2'})

# D2 - Quant Unit 3 (practice + review + cross-section)
rows_d.append({'id': 'D.5', 'case': 'Q-U3 Practice — QLE_2 + lesson failures',
    'precondition': 'QLE_2_ALG2 = Fail, PSS_Lesson_2 = Fail',
    'action': 'Load study plan, Quant Unit 3 Practice',
    'expected': f'Practice suggest box with redo: algebra_2, pss_lesson_2'})
rows_d.append({'id': 'D.6', 'case': 'Q-U3 Review — QLE_1 failures (prev unit)',
    'precondition': 'QLE_1_NP1 = Fail, QLE_1_PSS1 = Fail',
    'action': 'Load study plan, Quant Unit 3 Review',
    'expected': f'Review suggest box with redo: number_props_1, pss_lesson_1.\nText: "{REVIEW_SUGGEST}"'})
rows_d.append({'id': 'D.7', 'case': 'Q-U3 Review — No QLE_1 failures',
    'precondition': 'All QLE_1_* = Pass or no signals',
    'action': 'Load study plan', 'expected': 'NO review suggest box at Quant U3'})

# Cross-section: verbal_review_2
rows_d.append({'id': 'D.8', 'case': 'Q-U3 Cross — VR2 FAIL',
    'precondition': 'Unit2_Verbal_Review_Pass_or_Fail = "Fail"',
    'action': 'Load study plan, Quant Unit 3 Review',
    'expected': 'Cross-suggest: "Also suggested:" with links to rc_exercise_1, verbal_review_2'})
rows_d.append({'id': 'D.9', 'case': 'Q-U3 Cross — VR2 PASS',
    'precondition': 'Unit2_Verbal_Review_Pass_or_Fail = "Pass"',
    'action': 'Load study plan, Quant Unit 3 Review',
    'expected': 'Cross-suggest: "Also suggested:" with link to verbal_review_2 only'})
rows_d.append({'id': 'D.10', 'case': 'Q-U3 Cross — VR2 NONE',
    'precondition': 'No Unit2_Verbal_Review signal',
    'action': 'Load study plan', 'expected': 'NO cross-suggest at Quant U3'})

# D3 - Quant Unit 4
rows_d.append({'id': 'D.11', 'case': 'Q-U4 Practice — QLE_3 failures',
    'precondition': 'QLE_3_WP3 = Fail, QLE_3_WP4 = Fail',
    'action': 'Load study plan, Quant Unit 4 Practice',
    'expected': 'Practice suggest box with redo: word_problems_3, word_problems_4'})
rows_d.append({'id': 'D.12', 'case': 'Q-U4 Review — QLE_2 failures (prev unit)',
    'precondition': 'QLE_2_FPR1 = Fail',
    'action': 'Load study plan, Quant Unit 4 Review',
    'expected': 'Review suggest box with redo: fprs_1'})
rows_d.append({'id': 'D.13', 'case': 'Q-U4 Cross — VR3 FAIL',
    'precondition': 'Unit3_Verbal_Review_Pass_or_Fail = "Fail"',
    'action': 'Load study plan, Quant Unit 4 Review',
    'expected': 'Cross-suggest with links: cr_lesson_5, cr_exercise_4, verbal_review_3'})
rows_d.append({'id': 'D.14', 'case': 'Q-U4 Cross — VR3 PASS',
    'precondition': 'Unit3_Verbal_Review_Pass_or_Fail = "Pass"',
    'action': 'Load study plan', 'expected': 'Cross-suggest: verbal_review_3 only'})
rows_d.append({'id': 'D.15', 'case': 'Q-U4 Cross — VR3 NONE',
    'precondition': 'No Unit3_Verbal_Review signal',
    'action': 'Load study plan', 'expected': 'NO cross-suggest at Quant U4'})

# D4 - Quant Unit 5
rows_d.append({'id': 'D.16', 'case': 'Q-U5 Practice — QLE_4 failures',
    'precondition': 'QLE_4_WP5 = Fail',
    'action': 'Load study plan, Quant Unit 5 Practice',
    'expected': 'Practice suggest box with redo: word_problems_5'})
rows_d.append({'id': 'D.17', 'case': 'Q-U5 Review — QLE_3 failures (prev unit)',
    'precondition': 'QLE_3_ALG3 = Fail',
    'action': 'Load study plan, Quant Unit 5 Review',
    'expected': 'Review suggest box with redo: algebra_3'})
rows_d.append({'id': 'D.18', 'case': 'Q-U5 Cross — VR4 FAIL',
    'precondition': 'Unit4_Verbal_Review_Pass_or_Fail = "Fail"',
    'action': 'Load study plan, Quant Unit 5 Review',
    'expected': 'Cross-suggest with links: cr_lesson_6, cr_exercise_5, verbal_review_4'})
rows_d.append({'id': 'D.19', 'case': 'Q-U5 Cross — VR4 PASS',
    'precondition': 'Unit4_Verbal_Review = "Pass"',
    'action': 'Load study plan', 'expected': 'Cross-suggest: verbal_review_4 only'})
rows_d.append({'id': 'D.20', 'case': 'Q-U5 Cross — VR4 NONE',
    'precondition': 'No Unit4_Verbal_Review signal',
    'action': 'Load study plan', 'expected': 'NO cross-suggest at Quant U5'})

# D5 - Quant Unit 6
rows_d.append({'id': 'D.21', 'case': 'Q-U6 Practice — QLE_5 failures',
    'precondition': 'QLE_5_ALG5 = Fail, QLE_5_NP3 = Fail',
    'action': 'Load study plan, Quant Unit 6 Practice',
    'expected': 'Practice suggest box with redo: algebra_5, number_props_3'})
rows_d.append({'id': 'D.22', 'case': 'Q-U6 Review — QLE_4 failures (prev unit)',
    'precondition': 'QLE_4_ALG4 = Fail, QLE_4_WP6 = Fail',
    'action': 'Load study plan, Quant Unit 6 Review',
    'expected': 'Review suggest box with redo: algebra_4, word_problems_6'})
rows_d.append({'id': 'D.23', 'case': 'Q-U6 Cross — VR5 FAIL',
    'precondition': 'Unit5_Verbal_Review_Pass_or_Fail = "Fail"',
    'action': 'Load study plan, Quant Unit 6 Review',
    'expected': 'Cross-suggest with links: cr_lesson_7, cr_exercise_6, verbal_review_5'})
rows_d.append({'id': 'D.24', 'case': 'Q-U6 Cross — VR5 PASS',
    'precondition': 'Unit5_Verbal_Review = "Pass"',
    'action': 'Load study plan', 'expected': 'Cross-suggest: verbal_review_5 only'})
rows_d.append({'id': 'D.25', 'case': 'Q-U6 Cross — VR5 NONE',
    'precondition': 'No Unit5_Verbal_Review signal',
    'action': 'Load study plan', 'expected': 'NO cross-suggest at Quant U6'})

add_test_table(doc, rows_d)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION E - DI Suggestion Logic
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Section E - Data Insights Suggestion Logic', level=1)

rows_e = [
    {'id': 'E.1', 'case': 'DI U1 Review — QLE_5 failures',
     'precondition': 'QLE_5_WP7 = Fail, QLE_5_ALG5 = Fail',
     'action': 'Load verbal-first study plan, DI Unit 1 Review',
     'expected': f'Review suggest box with redo: word_problems_7, algebra_5.\nText: "{REVIEW_SUGGEST}"'},
    {'id': 'E.2', 'case': 'DI U1 Review — QLE_5 all pass',
     'precondition': 'All QLE_5_* = Pass',
     'action': 'Load study plan', 'expected': 'NO review suggest box at DI U1'},
    {'id': 'E.3', 'case': 'DI U1 Review — No QLE_5 signals',
     'precondition': 'No QLE_5_* in xAPI data',
     'action': 'Load study plan', 'expected': 'NO review suggest box at DI U1'},
    {'id': 'E.4', 'case': 'DI U2 — No suggestions',
     'precondition': 'Any xAPI state', 'action': 'Load study plan',
     'expected': 'DI Unit 2 has NO suggest boxes'},
    {'id': 'E.5', 'case': 'DI U3 — No suggestions',
     'precondition': 'Any xAPI state', 'action': 'Load study plan',
     'expected': 'DI Unit 3 has NO suggest boxes'},
]
add_test_table(doc, rows_e)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION F - Quant-First Path Differences
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Section F - Quant-First Path Differences', level=1)
doc.add_paragraph(
    'Tests specific to the quant-first path where cross-section suggest direction is REVERSED: '
    'verbal units suggest quant reviews (instead of quant units suggesting verbal reviews).'
)

rows_f = [
    # F1 - Verbal Unit 1 gets quant_review_6 + review suggest
    {'id': 'F.1', 'case': 'QF Verbal U1 — Review is quant_review_6',
     'precondition': 'User preference = quant',
     'action': 'Load quant-first study plan',
     'expected': 'Verbal Unit 1 Review = quant_review_6 (not empty)'},
    {'id': 'F.2', 'case': 'QF Verbal U1 — Review suggest from QLE_5',
     'precondition': 'QLE_5_NP3 = Fail',
     'action': 'Load quant-first study plan, Verbal U1 Review',
     'expected': f'Review suggest box with redo: number_props_3.\nText: "{REVIEW_SUGGEST}"'},
    {'id': 'F.3', 'case': 'QF Verbal U1 — No QLE_5 failures',
     'precondition': 'All QLE_5_* = Pass',
     'action': 'Load study plan', 'expected': 'NO review suggest box at QF Verbal U1'},

    # F2 - Verbal Unit 3: cross-suggest from quant_review_2
    {'id': 'F.4', 'case': 'QF V-U3 Cross — QR2 FAIL',
     'precondition': 'QRS_Unit2_ALG1 = Fail, QRS_Unit2_WP1 = Fail (quant_review_2 result = fail)',
     'action': 'Load quant-first plan, Verbal U3 Review',
     'expected': 'Cross-suggest with links: algebra_1, word_problems_1, quant_exercise_1, quant_review_2'},
    {'id': 'F.5', 'case': 'QF V-U3 Cross — QR2 PASS',
     'precondition': 'All QRS_Unit2_* = Pass (quant_review_2 result = pass)',
     'action': 'Load quant-first plan', 'expected': 'Cross-suggest: quant_review_2 only'},
    {'id': 'F.6', 'case': 'QF V-U3 Cross — QR2 NONE',
     'precondition': 'No QRS_Unit2_* signals',
     'action': 'Load quant-first plan', 'expected': 'NO cross-suggest at QF V-U3'},

    # F3 - Verbal Unit 4: cross-suggest from quant_review_3
    {'id': 'F.7', 'case': 'QF V-U4 Cross — QR3 FAIL',
     'precondition': 'QRS_Unit3_FPR1 = Fail (quant_review_3 result = fail)',
     'action': 'Load quant-first plan, Verbal U4 Review',
     'expected': 'Cross-suggest with links: fprs_1, quant_exercise_2, quant_review_3'},
    {'id': 'F.8', 'case': 'QF V-U4 Cross — QR3 PASS',
     'precondition': 'All QRS_Unit3_* = Pass',
     'action': 'Load quant-first plan', 'expected': 'Cross-suggest: quant_review_3 only'},
    {'id': 'F.9', 'case': 'QF V-U4 Cross — QR3 NONE',
     'precondition': 'No QRS_Unit3_* signals',
     'action': 'Load quant-first plan', 'expected': 'NO cross-suggest at QF V-U4'},

    # F4 - Verbal Unit 5: cross-suggest from quant_review_4
    {'id': 'F.10', 'case': 'QF V-U5 Cross — QR4 FAIL',
     'precondition': 'QRS_Unit4_WP3 = Fail (quant_review_4 result = fail)',
     'action': 'Load quant-first plan, Verbal U5 Review',
     'expected': 'Cross-suggest with links: word_problems_3, quant_exercise_3, quant_review_4'},
    {'id': 'F.11', 'case': 'QF V-U5 Cross — QR4 PASS',
     'precondition': 'All QRS_Unit4_* = Pass',
     'action': 'Load quant-first plan', 'expected': 'Cross-suggest: quant_review_4 only'},
    {'id': 'F.12', 'case': 'QF V-U5 Cross — QR4 NONE',
     'precondition': 'No QRS_Unit4_* signals',
     'action': 'Load quant-first plan', 'expected': 'NO cross-suggest at QF V-U5'},

    # F5 - Verbal Unit 6: cross-suggest from quant_review_5
    {'id': 'F.13', 'case': 'QF V-U6 Cross — QR5 FAIL',
     'precondition': 'QRS_Unit5_WP6 = Fail (quant_review_5 result = fail)',
     'action': 'Load quant-first plan, Verbal U6 Review',
     'expected': 'Cross-suggest with links: word_problems_6, quant_exercise_4, quant_review_5'},
    {'id': 'F.14', 'case': 'QF V-U6 Cross — QR5 PASS',
     'precondition': 'All QRS_Unit5_* = Pass',
     'action': 'Load quant-first plan', 'expected': 'Cross-suggest: quant_review_5 only'},
    {'id': 'F.15', 'case': 'QF V-U6 Cross — QR5 NONE',
     'precondition': 'No QRS_Unit5_* signals',
     'action': 'Load quant-first plan', 'expected': 'NO cross-suggest at QF V-U6'},

    # F6 - DI Unit 1 has NO review in quant-first (quant_review_6 moved to Verbal U1)
    {'id': 'F.16', 'case': 'QF DI U1 — No review slot',
     'precondition': 'User preference = quant',
     'action': 'Load quant-first study plan',
     'expected': 'DI Unit 1 Review = (none). No review suggest box.'},
]
add_test_table(doc, rows_f)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION G - Edge Cases & Data Integrity
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Section G - Edge Cases & Data Integrity', level=1)

rows_g = [
    {'id': 'G.1', 'case': 'Smart quotes in xAPI object name',
     'precondition': 'xAPI completed statement has object.definition.name containing curly quotes: \u201cCR_Exercise_4_Pass_or_Fail\u201d: \u201cFail\u201d',
     'action': 'Load study plan',
     'expected': 'Smart quotes sanitized to straight quotes. CR Exercise 4 correctly parsed as "Fail". Suggestion renders.'},
    {'id': 'G.2', 'case': 'Trailing whitespace in pass/fail value',
     'precondition': 'xAPI signal: CR_Exercise_5_Pass_or_Fail = "Pass " (trailing space)',
     'action': 'Load study plan',
     'expected': 'Value trimmed to "Pass". No false suggestion triggered.'},
    {'id': 'G.3', 'case': 'Empty result block in xAPI',
     'precondition': 'xAPI completed statement has result: {} (empty object)',
     'action': 'Load study plan',
     'expected': 'Lesson marked as "completed" (empty result block is valid completion).'},
    {'id': 'G.4', 'case': 'Fresh user — no xAPI data',
     'precondition': 'User has zero xAPI statements',
     'action': 'Load study plan',
     'expected': 'All lessons show "not-started". All pass/fail results are "none". No suggestion boxes appear.'},
    {'id': 'G.5', 'case': 'Mixed signals in same exercise',
     'precondition': 'QLE_1_ALG1 = Fail, QLE_1_NP1 = Pass, QLE_1_WP1 = Fail, QLE_1_PSS1 = Pass',
     'action': 'Load study plan, Quant Unit 2 Practice',
     'expected': 'Practice suggest shows redo for algebra_1 and word_problems_1 only (not number_props_1 or pss_lesson_1).'},
    {'id': 'G.6', 'case': 'Duplicate xAPI statements (most recent wins)',
     'precondition': 'Two completed statements for same activity: older = Fail, newer = Pass',
     'action': 'Load study plan',
     'expected': 'Most recent signal used (Pass). No false suggestion.'},
    {'id': 'G.7', 'case': 'BOM in xAPI object name',
     'precondition': 'xAPI object name starts with UTF-8 BOM (\\xEF\\xBB\\xBF)',
     'action': 'Load study plan',
     'expected': 'BOM stripped before JSON parsing. Pass/fail signals parsed correctly.'},
    {'id': 'G.8', 'case': 'Fallback regex when JSON decode fails',
     'precondition': 'xAPI object name is malformed JSON but matches: "VAR_Pass_or_Fail": "Fail"',
     'action': 'Load study plan',
     'expected': 'Regex fallback extracts variable and value. Suggestion renders correctly.'},
    {'id': 'G.9', 'case': 'xAPI slug not configured for DI lesson',
     'precondition': 'Admin has not set xAPI URL for di_lesson_4',
     'action': 'Load study plan',
     'expected': 'di_lesson_4 shows "not-started" (no tracking possible). No crash.'},
    {'id': 'G.10', 'case': 'result.completion = true with no pass/fail object',
     'precondition': 'xAPI completed statement with result.completion=true but no JSON in object name',
     'action': 'Load study plan',
     'expected': 'Lesson marked "completed" but no pass/fail signal extracted. Result = "none".'},
]
add_test_table(doc, rows_g)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# SECTION H - Suggest Box Rendering
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Section H - Suggest Box Rendering', level=1)
doc.add_paragraph(
    'Verify the visual rendering and HTML structure of suggestion boxes '
    'produced by gmat_sp_build_suggest_html().'
)

rows_h = [
    {'id': 'H.1', 'case': 'Practice suggest box position',
     'precondition': 'Quant U2 has QLE_1 failures',
     'action': 'Load study plan, scroll to Quant Unit 2',
     'expected': 'Orange suggest box (.gmat-sp-suggest) appears ABOVE the first practice lesson, BELOW the last learn lesson.'},
    {'id': 'H.2', 'case': 'Review suggest box position',
     'precondition': 'Quant U3 has QLE_1 failures',
     'action': 'Load study plan, scroll to Quant Unit 3',
     'expected': 'Review suggest box appears ABOVE the first review lesson, BELOW the last practice lesson.'},
    {'id': 'H.3', 'case': 'Suggest box text content',
     'precondition': 'CR_Exercise_4 = Fail',
     'action': 'Load study plan, Verbal Unit 4',
     'expected': f'Box contains exact text: "{SUGGEST_TEXT}"'},
    {'id': 'H.4', 'case': 'Redo links use correct names',
     'precondition': 'QLE_1_ALG1 = Fail',
     'action': 'Load study plan, check redo link',
     'expected': 'Link text = "Algebra Lesson 1" (from lesson label). Link href = "#sp-lesson-algebra_1".'},
    {'id': 'H.5', 'case': 'Cross-suggest prefix',
     'precondition': 'Unit2_Verbal_Review = Fail',
     'action': 'Load verbal-first plan, Quant U3 Review',
     'expected': 'Cross-suggest paragraph starts with bold "Also suggested:" followed by lesson links.'},
    {'id': 'H.6', 'case': 'Empty suggest — no box rendered',
     'precondition': 'All pass or no signals',
     'action': 'Load study plan',
     'expected': 'No .gmat-sp-suggest elements in DOM for units with no failures.'},
    {'id': 'H.7', 'case': 'Suggest box has orange icon',
     'precondition': 'Any failure triggering a suggest box',
     'action': 'Inspect suggest box HTML',
     'expected': 'Box contains SVG icon in .gmat-sp-suggest__icon. Header text = "Suggested areas of focus".'},
    {'id': 'H.8', 'case': 'Multiple redo links separated correctly',
     'precondition': 'QLE_1_ALG1 = Fail, QLE_1_WP1 = Fail, QLE_1_NP1 = Fail',
     'action': 'Load study plan, Quant U2 Practice suggest',
     'expected': 'Redo links appear as: "Recommend redoing: Algebra Lesson 1, Word Problems Lesson 1, Number Properties Lesson 1" with comma-separated links.'},
    {'id': 'H.9', 'case': 'Suggest links use ampersand separator',
     'precondition': 'Cross-suggest fail with multiple links',
     'action': 'Inspect cross-suggest HTML',
     'expected': 'Suggest links separated by " & " (ampersand with spaces).'},
    {'id': 'H.10', 'case': 'Responsive layout on mobile (< 768px)',
     'precondition': 'Any failure triggering suggest box',
     'action': 'Resize viewport to 375px width',
     'expected': 'Suggest box spans full width, text wraps properly, links are tappable.'},
]
add_test_table(doc, rows_h)

# ── Save ─────────────────────────────────────────────────────────────

output_path = 'GMAT_Study_Plan_Test_Cases.docx'
doc.save(output_path)

# Count test cases
total = (len(rows_a1) + len(rows_a2) + len(rows_b1) + len(rows_b2) +
         len(rows_b3) + len(rows_b4) + len(rows_b5) + len(rows_c) +
         len(rows_d) + len(rows_e) + len(rows_f) + len(rows_g) + len(rows_h))

print(f'Generated: {output_path}')
print(f'Total test cases: {total}')
print('Sections: A (Structure), B (xAPI Functions), C (Verbal Suggest), '
      'D (Quant Suggest), E (DI Suggest), F (Quant-First), G (Edge Cases), H (Rendering)')
